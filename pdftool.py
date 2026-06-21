import io
import os
import shutil
import sys
import uuid
import zipfile
from pathlib import Path
from PIL import Image
from pypdf import PdfReader, PdfWriter
import pypdfium2 as pdfium
from pdf2docx import Converter as _PdfDocxConverter
from docx import Document as _DocxDocument
from docx.shared import Pt as _Pt
import openpyxl
from reportlab.lib.pagesizes import A4 as _A4, LETTER as _LETTER
from reportlab.lib.styles import getSampleStyleSheet as _getSampleStyleSheet
from reportlab.lib.units import cm as _cm
from reportlab.platypus import SimpleDocTemplate as _SimpleDocTemplate, Paragraph as _Paragraph, Spacer as _Spacer, Table as _RLTable, TableStyle as _TableStyle, PageBreak as _PageBreak
from reportlab.lib import colors as _rl_colors

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif", ".gif"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx", ".doc"}
XLSX_EXTENSIONS = {".xlsx", ".xls"}


def _to_rgb(img):
    if img.mode == "RGB":
        return img
    if img.mode in ("RGBA", "LA"):
        bg = Image.new("RGB", img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[-1])
        return bg
    return img.convert("RGB")


def _resize(img, max_dimension):
    if max_dimension and (img.width > max_dimension or img.height > max_dimension):
        img.thumbnail((max_dimension, max_dimension), Image.LANCZOS)
    return img


def _save_image(img, output_path, quality=60, fmt=None):
    if fmt is None:
        fmt = Path(output_path).suffix.lstrip(".").upper() or "PNG"
    fmt = fmt.upper()
    if fmt in ("JPG", "JPEG"):
        img = _to_rgb(img)
        img.save(output_path, "JPEG", quality=quality, optimize=True, progressive=True)
    elif fmt == "PNG":
        img.save(output_path, "PNG", optimize=True)
    elif fmt == "WEBP":
        img.save(output_path, "WEBP", quality=quality, method=6)
    elif fmt == "TIFF":
        img.save(output_path, "TIFF", compression="tiff_lzw")
    else:
        if fmt in ("BMP",):
            img = _to_rgb(img)
        img.save(output_path, fmt, optimize=True)


def merge_pdfs(input_paths, output_path):
    pdf_paths = []
    image_paths = []
    for path in input_paths:
        ext = Path(path).suffix.lower()
        if ext in IMAGE_EXTENSIONS:
            image_paths.append(path)
        elif ext in PDF_EXTENSIONS:
            pdf_paths.append(path)
        else:
            raise ValueError(f"unsupported input file: {path}")
    writer = PdfWriter()
    if image_paths:
        tmp = output_path + f".imgs_{uuid.uuid4().hex}.pdf"
        try:
            images_to_pdf(image_paths, tmp)
            for page in PdfReader(tmp).pages:
                writer.add_page(page)
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)
    for path in pdf_paths:
        reader = PdfReader(path)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                raise ValueError(f"encrypted PDF not supported: {path}")
        for page in reader.pages:
            writer.add_page(page)
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_path


def images_to_pdf(input_paths, output_path):
    images = []
    for p in input_paths:
        img = Image.open(p)
        img.load()
        img = _to_rgb(img)
        images.append(img)
    if not images:
        raise ValueError("no images provided")
    first, *rest = images
    first.save(output_path, "PDF", save_all=True, append_images=rest, resolution=150.0)
    return output_path


def pdf_to_images(input_path, output_dir, fmt="png", dpi=150, quality=85):
    pdf = pdfium.PdfDocument(input_path)
    scale = dpi / 72
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    paths = []
    for i in range(len(pdf)):
        page = pdf[i]
        bitmap = page.render(scale=scale)
        pil_image = bitmap.to_pil()
        out_path = output_dir / f"page_{i + 1:04d}.{fmt}"
        _save_image(pil_image, str(out_path), quality=quality, fmt=fmt)
        paths.append(str(out_path))
    return paths


def compress_image(input_path, output_path, quality=60, max_dimension=None):
    img = Image.open(input_path)
    img.load()
    img = _resize(img, max_dimension)
    _save_image(img, output_path, quality=quality)
    return output_path


def compress_pdf(input_path, output_path, quality=60, max_dimension=1600, dpi=120, mode="auto"):
    if mode == "rasterize":
        return _compress_pdf_rasterize(input_path, output_path, quality, max_dimension, dpi)
    if mode == "native":
        return _compress_pdf_native(input_path, output_path, quality, max_dimension)
    try:
        return _compress_pdf_native(input_path, output_path, quality, max_dimension)
    except Exception:
        return _compress_pdf_rasterize(input_path, output_path, quality, max_dimension, dpi)


def _compress_pdf_native(input_path, output_path, quality=60, max_dimension=1600):
    reader = PdfReader(input_path)
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:
            raise ValueError("encrypted PDF not supported")
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    for page in writer.pages:
        try:
            for img in page.images:
                try:
                    pil_img = img.image
                    pil_img = _to_rgb(pil_img)
                    pil_img = _resize(pil_img, max_dimension)
                    img.replace(pil_img, quality=quality)
                except Exception:
                    continue
        except Exception:
            pass
        try:
            page.compress_content_streams()
        except Exception:
            pass
        try:
            for fn in page.inline_images:
                pass
        except Exception:
            pass
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_path


def _compress_pdf_rasterize(input_path, output_path, quality=60, max_dimension=1600, dpi=120):
    pdf = pdfium.PdfDocument(input_path)
    scale = dpi / 72
    images = []
    for i in range(len(pdf)):
        page = pdf[i]
        bitmap = page.render(scale=scale)
        pil_image = bitmap.to_pil()
        pil_image = _to_rgb(pil_image)
        pil_image = _resize(pil_image, max_dimension)
        images.append(pil_image)
    if not images:
        raise ValueError("PDF has no pages")
    first, *rest = images
    tmp_img_dir = output_path + f".imgs_{uuid.uuid4().hex}"
    try:
        Path(tmp_img_dir).mkdir(parents=True, exist_ok=True)
        tmp_paths = []
        for idx, im in enumerate(images, start=1):
            tp = os.path.join(tmp_img_dir, f"p_{idx:04d}.jpg")
            _save_image(im, tp, quality=quality, fmt="jpg")
            tmp_paths.append(tp)
        imgs = [Image.open(p) for p in tmp_paths]
        imgs[0].save(output_path, "PDF", save_all=True, append_images=imgs[1:], resolution=72.0)
    finally:
        shutil.rmtree(tmp_img_dir, ignore_errors=True)
    return output_path


def convert_image(input_path, output_path, quality=90):
    img = Image.open(input_path)
    img.load()
    _save_image(img, output_path, quality=quality)
    return output_path


def convert_file(input_path, output_path, quality=85, dpi=150):
    in_ext = Path(input_path).suffix.lower()
    out_ext = Path(output_path).suffix.lower()
    out_path = Path(output_path)
    if in_ext == ".pdf" and out_ext in IMAGE_EXTENSIONS:
        out_dir = out_path.parent if str(out_path.parent) else Path(".")
        fmt = out_ext.lstrip(".")
        paths = pdf_to_images(input_path, str(out_dir), fmt=fmt, dpi=dpi, quality=quality)
        if len(paths) == 1:
            os.replace(paths[0], output_path)
            return [output_path]
        return paths
    if in_ext in IMAGE_EXTENSIONS and out_ext == ".pdf":
        images_to_pdf([input_path], output_path)
        return [output_path]
    if in_ext in IMAGE_EXTENSIONS and out_ext in IMAGE_EXTENSIONS:
        convert_image(input_path, output_path, quality=quality)
        return [output_path]
    if in_ext == ".pdf" and out_ext in DOCX_EXTENSIONS:
        pdf_to_docx(input_path, output_path)
        return [output_path]
    if in_ext in DOCX_EXTENSIONS and out_ext == ".pdf":
        docx_to_pdf(input_path, output_path)
        return [output_path]
    if in_ext in XLSX_EXTENSIONS and out_ext == ".pdf":
        xlsx_to_pdf(input_path, output_path)
        return [output_path]
    raise ValueError(f"unsupported conversion: {in_ext} -> {out_ext}")


def make_zip(paths_or_dir, zip_path):
    paths = []
    if isinstance(paths_or_dir, (list, tuple)):
        paths = list(paths_or_dir)
    else:
        p = Path(paths_or_dir)
        if p.is_dir():
            paths = sorted([str(x) for x in p.iterdir() if x.is_file()])
        else:
            paths = [str(p)]
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for fp in paths:
            z.write(fp, Path(fp).name)
    return zip_path


def build_parser():
    import argparse
    parser = argparse.ArgumentParser(description="PDF merger, compressor and file converter")
    subparsers = parser.add_subparsers(dest="command", required=True)
    merge_parser = subparsers.add_parser("merge")
    merge_parser.add_argument("inputs", nargs="+")
    merge_parser.add_argument("-o", "--output", required=True)
    compress_parser = subparsers.add_parser("compress")
    compress_parser.add_argument("input")
    compress_parser.add_argument("-o", "--output", required=True)
    compress_parser.add_argument("-q", "--quality", type=int, default=60)
    compress_parser.add_argument("-m", "--max-dimension", type=int, default=1600)
    compress_parser.add_argument("--dpi", type=int, default=120)
    compress_parser.add_argument("--mode", choices=["auto", "native", "rasterize"], default="auto")
    convert_parser = subparsers.add_parser("convert")
    convert_parser.add_argument("input")
    convert_parser.add_argument("-o", "--output", required=True)
    convert_parser.add_argument("-q", "--quality", type=int, default=85)
    convert_parser.add_argument("--dpi", type=int, default=150)
    return parser


def main():
    parser = build_parser()
    args = parser.parse_args()
    if args.command == "merge":
        merge_pdfs(args.inputs, args.output)
        print(f"merged {len(args.inputs)} files into {args.output}")
    elif args.command == "compress":
        ext = Path(args.input).suffix.lower()
        if ext == ".pdf":
            compress_pdf(args.input, args.output, quality=args.quality, max_dimension=args.max_dimension, dpi=args.dpi, mode=args.mode)
        elif ext in IMAGE_EXTENSIONS:
            compress_image(args.input, args.output, quality=args.quality, max_dimension=args.max_dimension)
        else:
            print(f"unsupported file type for compression: {ext}")
            sys.exit(1)
        before = os.path.getsize(args.input)
        after = os.path.getsize(args.output)
        ratio = (1 - after / before) * 100 if before else 0
        print(f"{args.input}: {before} bytes -> {args.output}: {after} bytes ({ratio:.1f}% smaller)")
    elif args.command == "convert":
        paths = convert_file(args.input, args.output, quality=args.quality, dpi=args.dpi)
        print(f"converted {args.input} -> {len(paths)} output file(s)")


def pdf_to_docx(input_path, output_path):
    cv = _PdfDocxConverter(input_path)
    try:
        cv.convert(output_path)
    finally:
        cv.close()
    return output_path


def docx_to_pdf(input_path, output_path):
    doc = _DocxDocument(input_path)
    styles = _getSampleStyleSheet()
    body_style = styles["BodyText"]
    body_style.fontSize = 11
    body_style.leading = 15
    body_style.spaceAfter = 6
    h_style = styles["Heading1"]
    h_style.fontSize = 16
    h_style.leading = 20
    h_style.spaceAfter = 10
    h2_style = styles["Heading2"]
    h2_style.fontSize = 13
    h2_style.leading = 17
    h2_style.spaceAfter = 8
    doc_pdf = _SimpleDocTemplate(
        str(output_path),
        pagesize=_A4,
        leftMargin=2 * _cm,
        rightMargin=2 * _cm,
        topMargin=2 * _cm,
        bottomMargin=2 * _cm,
        title="Converted Document",
    )
    story = []
    for para in doc.paragraphs:
        text = para.text or ""
        if not text.strip():
            story.append(_Spacer(1, 0.2 * _cm))
            continue
        style_name = (para.style.name or "").lower() if para.style else ""
        if "heading 1" in style_name or "title" in style_name:
            story.append(_Paragraph(_xml_escape(text), h_style))
        elif "heading" in style_name:
            story.append(_Paragraph(_xml_escape(text), h2_style))
        else:
            bold = any(run.bold for run in para.runs if run.bold is not None)
            italic = any(run.italic for run in para.runs if run.italic is not None)
            inline_style = body_style
            escaped = _xml_escape(text)
            if bold and italic:
                escaped = f"<b><i>{escaped}</i></b>"
            elif bold:
                escaped = f"<b>{escaped}</b>"
            elif italic:
                escaped = f"<i>{escaped}</i>"
            story.append(_Paragraph(escaped, inline_style))
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            rows_data.append([_xml_escape(cell.text or "") for cell in row.cells])
        if not rows_data:
            continue
        n_cols = max(len(r) for r in rows_data)
        for r in rows_data:
            while len(r) < n_cols:
                r.append("")
        try:
            tbl = _RLTable(rows_data, repeatRows=1)
            tbl.setStyle(_TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), _rl_colors.HexColor("#5a67d8")),
                ("TEXTCOLOR", (0, 0), (-1, 0), _rl_colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("GRID", (0, 0), (-1, -1), 0.5, _rl_colors.HexColor("#cbd5e0")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(tbl)
            story.append(_Spacer(1, 0.4 * _cm))
        except Exception:
            pass
    if not story:
        story.append(_Paragraph("(empty document)", body_style))
    doc_pdf.build(story)
    return output_path


def xlsx_to_pdf(input_path, output_path):
    wb = openpyxl.load_workbook(input_path, data_only=True)
    doc_pdf = _SimpleDocTemplate(
        str(output_path),
        pagesize=_A4,
        leftMargin=1.5 * _cm,
        rightMargin=1.5 * _cm,
        topMargin=1.5 * _cm,
        bottomMargin=1.5 * _cm,
        landscape=True,
    )
    styles = _getSampleStyleSheet()
    title_style = styles["Title"]
    title_style.fontSize = 16
    story = []
    for ws in wb.worksheets:
        story.append(_Paragraph(_xml_escape(ws.title or "Sheet"), title_style))
        story.append(_Spacer(1, 0.3 * _cm))
        rows_data = []
        max_col = 0
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if not any(c.strip() for c in cells):
                continue
            rows_data.append([_xml_escape(c) for c in cells])
            if len(cells) > max_col:
                max_col = len(cells)
        if not rows_data:
            story.append(_Paragraph("(empty sheet)", styles["Italic"]))
            story.append(_PageBreak())
            continue
        if max_col == 0:
            max_col = 1
        col_widths = [_doc_pdf_available_width() / max_col] * max_col
        try:
            tbl = _RLTable(rows_data, colWidths=col_widths, repeatRows=1)
            tbl.setStyle(_TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), _rl_colors.HexColor("#5a67d8")),
                ("TEXTCOLOR", (0, 0), (-1, 0), _rl_colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("GRID", (0, 0), (-1, -1), 0.4, _rl_colors.HexColor("#cbd5e0")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(tbl)
        except Exception:
            pass
        story.append(_PageBreak())
    if not wb.worksheets:
        story.append(_Paragraph("(workbook has no sheets)", styles["BodyText"]))
    doc_pdf.build(story)
    return output_path


_doc_pdf_page_size = _A4
_doc_pdf_page_w, _doc_pdf_page_h = _A4


def _doc_pdf_available_width():
    return _doc_pdf_page_w - 3 * _cm


def _xml_escape(s):
    return (s.replace("&", "&amp;")
             .replace("<", "&lt;")
             .replace(">", "&gt;")
             .replace('"', "&quot;")
             .replace("'", "&#39;"))


if __name__ == "__main__":
    main()

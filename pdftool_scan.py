"""
Doc Scanner backend module
==========================
Premium document scanner functions:
- Auto edge detection & perspective correction
- Quality filters (Original, B&W, Grayscale, Magic Color, Enhanced)
- Manual crop with coordinates
- Rotation
- Brightness/contrast/sharpness adjustment
- OCR text extraction (optional, when Tesseract is available)
- Multi-page scan -> PDF -> DOCX
"""

import io
import os
import math
import uuid
import shutil
import zipfile
from pathlib import Path

import cv2
import numpy as np
from PIL import Image, ImageEnhance, ImageFilter, ImageOps

# OCR is optional — only works if the tesseract BINARY is installed.
# Importing the pytesseract package always succeeds once it's pip-installed,
# even if the underlying `tesseract` executable is missing, so we explicitly
# probe the binary itself here. Render's native Python environment, for
# example, does not allow apt-get/sudo and has no tesseract preinstalled —
# without this check, ocr_available() would incorrectly report True there.
try:
    import pytesseract
    pytesseract.get_tesseract_version()
    TESSERACT_AVAILABLE = True
except Exception:
    pytesseract = None
    TESSERACT_AVAILABLE = False


# ─────────────────────────────────────────────────────────────
# Image helpers
# ─────────────────────────────────────────────────────────────

def _pil_to_cv(img):
    """PIL RGB -> OpenCV BGR"""
    rgb = np.array(img.convert("RGB"))
    return cv2.cvtColor(rgb, cv2.COLOR_RGB2BGR)


def _cv_to_pil(arr):
    """OpenCV BGR -> PIL RGB"""
    rgb = cv2.cvtColor(arr, cv2.COLOR_BGR2RGB)
    return Image.fromarray(rgb)


def _load_image(path_or_bytes):
    if isinstance(path_or_bytes, (bytes, bytearray)):
        arr = np.frombuffer(path_or_bytes, dtype=np.uint8)
        bgr = cv2.imdecode(arr, cv2.IMREAD_COLOR)
    else:
        bgr = cv2.imread(str(path_or_bytes), cv2.IMREAD_COLOR)
    if bgr is None:
        raise ValueError("could not decode image")
    return bgr


def _save_pil(img, path, fmt=None, quality=92):
    if fmt is None:
        fmt = Path(path).suffix.lstrip(".").upper() or "PNG"
    fmt = fmt.upper()
    if fmt in ("JPG", "JPEG"):
        img = img.convert("RGB")
        img.save(path, "JPEG", quality=quality, optimize=True)
    elif fmt == "PNG":
        img.save(path, "PNG", optimize=True)
    elif fmt == "WEBP":
        img.save(path, "WEBP", quality=quality, method=6)
    else:
        img.save(path, fmt)


# ─────────────────────────────────────────────────────────────
# Auto-edge detection + perspective correction (premium feature)
# ─────────────────────────────────────────────────────────────

def _order_points(pts):
    """Order 4 points as TL, TR, BR, BL."""
    rect = np.zeros((4, 2), dtype="float32")
    s = pts.sum(axis=1)
    rect[0] = pts[np.argmin(s)]      # top-left (smallest sum)
    rect[2] = pts[np.argmax(s)]      # bottom-right (largest sum)
    diff = np.diff(pts, axis=1)
    rect[1] = pts[np.argmin(diff)]   # top-right (smallest diff)
    rect[3] = pts[np.argmax(diff)]   # bottom-left (largest diff)
    return rect


def _four_point_transform(image, pts):
    """Apply perspective warp given 4 points."""
    rect = _order_points(pts)
    (tl, tr, br, bl) = rect
    width_a = np.linalg.norm(br - bl)
    width_b = np.linalg.norm(tr - tl)
    height_a = np.linalg.norm(tr - br)
    height_b = np.linalg.norm(tl - bl)
    max_w = int(max(width_a, width_b))
    max_h = int(max(height_a, height_b))
    dst = np.array([
        [0, 0],
        [max_w - 1, 0],
        [max_w - 1, max_h - 1],
        [0, max_h - 1]
    ], dtype="float32")
    M = cv2.getPerspectiveTransform(rect, dst)
    warped = cv2.warpPerspective(image, M, (max_w, max_h))
    return warped


def _find_document_contour(image):
    """Find the most-likely document quadrilateral in an image."""
    h, w = image.shape[:2]
    # resize for faster processing
    scale = 800.0 / max(h, w) if max(h, w) > 800 else 1.0
    if scale < 1.0:
        small = cv2.resize(image, None, fx=scale, fy=scale)
    else:
        small = image

    gray = cv2.cvtColor(small, cv2.COLOR_BGR2GRAY)
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)
    edges = cv2.Canny(blurred, 50, 150)
    edges = cv2.dilate(edges, np.ones((3, 3), np.uint8), iterations=1)

    contours, _ = cv2.findContours(edges.copy(), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    contours = sorted(contours, key=cv2.contourArea, reverse=True)[:8]

    doc = None
    img_area = small.shape[0] * small.shape[1]
    for c in contours:
        if cv2.contourArea(c) < img_area * 0.15:
            continue
        peri = cv2.arcLength(c, True)
        approx = cv2.approxPolyDP(c, 0.02 * peri, True)
        if len(approx) == 4:
            doc = approx.reshape(4, 2) / scale
            break

    if doc is None:
        # fallback: use the whole image
        doc = np.array([[0, 0], [w - 1, 0], [w - 1, h - 1], [0, h - 1]], dtype="float32")
    return doc


def auto_perspective_correct(image_bgr):
    """Detect document edges and apply perspective correction."""
    try:
        pts = _find_document_contour(image_bgr)
        return _four_point_transform(image_bgr, pts)
    except Exception:
        return image_bgr


# ─────────────────────────────────────────────────────────────
# Filters / enhancements
# ─────────────────────────────────────────────────────────────

def apply_filter(image_bgr, filter_name="magic_color", brightness=1.0,
                 contrast=1.0, sharpness=1.0):
    """
    Apply a named filter. Returns BGR numpy array.

    filter_name ∈ {original, bw, grayscale, magic_color, enhanced, sharpen}
    """
    img = image_bgr

    if filter_name == "bw":
        # Adaptive threshold (best for documents)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        # light denoise first
        gray = cv2.bilateralFilter(gray, 9, 75, 75)
        thr = cv2.adaptiveThreshold(
            gray, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            31, 15
        )
        img = cv2.cvtColor(thr, cv2.COLOR_GRAY2BGR)

    elif filter_name == "grayscale":
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        img = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)

    elif filter_name == "magic_color":
        # Adobe-style Magic Color: punchy contrast, light denoise, mild saturation
        pil = _cv_to_pil(img)
        # 1) light denoise
        pil = pil.filter(ImageFilter.SMOOTH)
        # 2) bump contrast
        pil = ImageEnhance.Contrast(pil).enhance(1.45)
        # 3) boost saturation a touch
        pil = ImageEnhance.Color(pil).enhance(1.18)
        # 4) white balance toward bright background
        pil = ImageEnhance.Brightness(pil).enhance(1.08)
        img = _pil_to_cv(pil)
        # 5) sharpen
        kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]], dtype=np.float32)
        img = cv2.filter2D(img, -1, kernel)

    elif filter_name == "enhanced":
        # Generic "enhanced" — clarity boost without going monochrome
        pil = _cv_to_pil(img)
        pil = ImageEnhance.Contrast(pil).enhance(1.25)
        pil = ImageEnhance.Sharpness(pil).enhance(1.6)
        pil = ImageEnhance.Color(pil).enhance(1.05)
        img = _pil_to_cv(pil)

    elif filter_name == "sharpen":
        kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]], dtype=np.float32)
        img = cv2.filter2D(img, -1, kernel)

    # else "original": no-op

    # Per-knob adjustments on top of filter
    if brightness != 1.0 or contrast != 1.0 or sharpness != 1.0:
        pil = _cv_to_pil(img)
        if brightness != 1.0:
            pil = ImageEnhance.Brightness(pil).enhance(brightness)
        if contrast != 1.0:
            pil = ImageEnhance.Contrast(pil).enhance(contrast)
        if sharpness != 1.0:
            pil = ImageEnhance.Sharpness(pil).enhance(sharpness)
        img = _pil_to_cv(pil)

    return img


# ─────────────────────────────────────────────────────────────
# Crop / rotate
# ─────────────────────────────────────────────────────────────

def crop_image(image_bgr, x, y, w, h):
    """Crop by pixel rectangle. Returns BGR."""
    H, W = image_bgr.shape[:2]
    x = max(0, int(x)); y = max(0, int(y))
    w = min(int(w), W - x); h = min(int(h), H - y)
    if w <= 0 or h <= 0:
        return image_bgr
    return image_bgr[y:y + h, x:x + w].copy()


def rotate_image(image_bgr, angle):
    """Rotate around center, expanding canvas to fit. Returns BGR."""
    if not angle:
        return image_bgr
    angle = float(angle) % 360
    h, w = image_bgr.shape[:2]
    M = cv2.getRotationMatrix2D((w / 2, h / 2), angle, 1.0)
    cos = abs(M[0, 0]); sin = abs(M[0, 1])
    nw = int(h * sin + w * cos)
    nh = int(h * cos + w * sin)
    M[0, 2] += nw / 2 - w / 2
    M[1, 2] += nh / 2 - h / 2
    return cv2.warpAffine(image_bgr, M, (nw, nh),
                          flags=cv2.INTER_CUBIC,
                          borderMode=cv2.BORDER_REPLICATE)


# ─────────────────────────────────────────────────────────────
# Pipeline: apply all settings in one shot
# ─────────────────────────────────────────────────────────────

def process_scan(input_path, output_path,
                 perspective=False,
                 filter_name="magic_color",
                 brightness=1.0, contrast=1.0, sharpness=1.0,
                 rotate=0,
                 crop=None):
    """
    Full scan pipeline.
    crop: dict {x,y,w,h} in original-image pixels or None

    Order of operations: crop -> rotate -> perspective -> filter.
    Crop runs FIRST because the frontend's crop UI always measures
    x/y/w/h against the as-uploaded photo (it crops on whatever image
    is currently displayed, which is the original until a page is
    exported) — running crop first keeps those coordinates valid
    regardless of whether rotate/perspective are also requested,
    instead of requiring the crop box to already be in rotated or
    perspective-warped coordinate space.
    """
    img = _load_image(input_path)

    # 1) crop first (coordinates are relative to the original photo)
    if crop:
        img = crop_image(img, crop["x"], crop["y"], crop["w"], crop["h"])

    # 2) rotate
    if rotate:
        img = rotate_image(img, rotate)

    # 3) perspective correct (if requested)
    if perspective:
        img = auto_perspective_correct(img)

    # 4) filter + tweaks
    img = apply_filter(img,
                       filter_name=filter_name,
                       brightness=brightness,
                       contrast=contrast,
                       sharpness=sharpness)

    # 5) write
    out_dir = Path(output_path).parent
    out_dir.mkdir(parents=True, exist_ok=True)
    fmt = Path(output_path).suffix.lstrip(".").lower() or "jpg"
    if fmt in ("jpg", "jpeg"):
        cv2.imwrite(str(output_path), img,
                    [cv2.IMWRITE_JPEG_QUALITY, 92])
    elif fmt == "png":
        cv2.imwrite(str(output_path), img)
    elif fmt == "webp":
        cv2.imwrite(str(output_path), img,
                    [cv2.IMWRITE_WEBP_QUALITY, 92])
    else:
        cv2.imwrite(str(output_path), img)
    return output_path


# ─────────────────────────────────────────────────────────────
# OCR (optional)
# ─────────────────────────────────────────────────────────────

def ocr_image(image_path, lang="eng"):
    """Run OCR on a single image. Returns extracted text (may be empty)."""
    if not TESSERACT_AVAILABLE:
        return ""
    try:
        return pytesseract.image_to_string(str(image_path), lang=lang) or ""
    except Exception:
        return ""


def ocr_available():
    return TESSERACT_AVAILABLE


# ─────────────────────────────────────────────────────────────
# Multi-page scan -> PDF / DOCX
# ─────────────────────────────────────────────────────────────

def images_to_pdf_simple(image_paths, output_path):
    """Just stitch images into a PDF (no OCR)."""
    imgs = []
    for p in image_paths:
        im = Image.open(p)
        im.load()
        if im.mode != "RGB":
            im = im.convert("RGB")
        imgs.append(im)
    if not imgs:
        raise ValueError("no images provided")
    first, *rest = imgs
    first.save(str(output_path), "PDF",
               save_all=True, append_images=rest, resolution=200.0)
    return output_path


def images_to_docx(image_paths, output_path, ocr_lang="eng", use_ocr=True):
    """
    Pipeline:
      1) Build a PDF from the processed images.
      2) If OCR is enabled and tesseract is available, build a docx from the OCR text
         (one section per page, text + page heading). Otherwise, convert PDF -> docx
         using pdf2docx (image-only pages still get a working .docx).
    """
    # Step 1: build PDF
    tmp_pdf = str(output_path) + ".intermediate.pdf"
    images_to_pdf_simple(image_paths, tmp_pdf)

    if use_ocr and TESSERACT_AVAILABLE:
        # Build a real text docx with OCR'd content per page
        from docx import Document
        from docx.shared import Pt, Inches
        doc = Document()
        # Page-size A4
        for sec in doc.sections:
            sec.top_margin = Inches(0.7)
            sec.bottom_margin = Inches(0.7)
            sec.left_margin = Inches(0.7)
            sec.right_margin = Inches(0.7)

        for idx, p in enumerate(image_paths, start=1):
            heading = doc.add_paragraph()
            run = heading.add_run(f"Page {idx}")
            run.bold = True
            run.font.size = Pt(14)

            text = ocr_image(p, lang=ocr_lang) or ""
            if text.strip():
                for line in text.splitlines():
                    para = doc.add_paragraph()
                    r = para.add_run(line)
                    r.font.size = Pt(11)
            else:
                note = doc.add_paragraph()
                nr = note.add_run("(no text recognized on this page)")
                nr.italic = True
                nr.font.size = Pt(10)

            # Page break between pages (not after the last one)
            if idx < len(image_paths):
                doc.add_page_break()

        doc.save(str(output_path))
    else:
        # Fallback: PDF -> DOCX (image-based, still a valid docx)
        from pdf2docx import Converter
        cv = Converter(tmp_pdf)
        try:
            cv.convert(str(output_path))
        finally:
            cv.close()

    # cleanup intermediate pdf
    try:
        os.remove(tmp_pdf)
    except OSError:
        pass

    return output_path
